# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path


def set_cn_font(run, name='宋体', size=12, bold=False):
    run.bold = bold
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run(text)
    set_cn_font(r, name='黑体', size=16, bold=True)


def add_h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_cn_font(r, name='黑体', size=14, bold=True)


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5


def set_table_three_line(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)

    tblBorders = OxmlElement('w:tblBorders')
    for edge, val in [
        ('top', 'single'),
        ('left', 'nil'),
        ('bottom', 'single'),
        ('right', 'nil'),
        ('insideH', 'nil'),
        ('insideV', 'nil')
    ]:
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), val)
        if val != 'nil':
            elem.set(qn('w:sz'), '8')
            elem.set(qn('w:space'), '0')
            elem.set(qn('w:color'), '000000')
        tblBorders.append(elem)
    tblPr.append(tblBorders)

    # 表头下横线
    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
        bottom = tcBorders.find(qn('w:bottom'))
        if bottom is None:
            bottom = OxmlElement('w:bottom')
            tcBorders.append(bottom)
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '8')
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), '000000')


doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.font.size = Pt(12)

add_title(doc, '校园外卖实时调度系统开题报告（工程实践类）')
doc.add_paragraph('')

add_h1(doc, '一、课题的目的和意义')
add_para(doc, '本课题面向校园外卖配送场景，目标是开发一套可落地运行的实时调度系统，解决订单高峰期派单效率低、骑手路径不合理、配送过程不透明等实际问题。课题属于工程实践类，重点不在理论证明，而在系统设计、功能实现、联调测试和效果验证。')
add_para(doc, '通过本课题建设，系统可实现用户端下单与跟踪、骑手端接单与配送、管理端监控与调度，形成完整业务闭环。对学校后勤场景或类似即时配送业务具有较强实践参考价值。')

add_h1(doc, '二、课题的主要内容')
add_para(doc, '本课题围绕“下单—派单—配送—完成—统计”主流程开展，主要内容包括：')
add_para(doc, '（1）三端业务功能建设：完成用户端、骑手端、管理端的核心功能。')
add_para(doc, '（2）调度模块实现：设计规则派单策略，并结合贪婪思想生成配送顺序。')
add_para(doc, '（3）校园节点地图构建：抽象校园节点与路径关系，支持可视化展示。')
add_para(doc, '（4）实时状态同步：实现订单状态流转与消息推送机制。')
add_para(doc, '（5）系统联调与测试：开展功能测试、接口联调、场景测试和结果分析。')

add_h1(doc, '三、课题拟解决的主要问题')
add_para(doc, '（1）高峰时段订单并发下的派单效率问题：要求新订单能在短时间内完成自动分配。')
add_para(doc, '（2）骑手多单配送顺序问题：在实际道路约束下给出可执行的推荐路径，减少绕行。')
add_para(doc, '（3）三端状态一致性问题：用户端、骑手端、管理端对订单状态需实时一致。')
add_para(doc, '（4）调度过程可解释性问题：通过地图与轨迹展示提升系统可视化与可维护性。')
add_para(doc, '（5）工程可交付性问题：确保系统可运行、可演示、可测试、可复现。')

add_h1(doc, '四、采用的技术方案和实现方法')
add_para(doc, '技术方案采用前后端分离架构。后端负责业务规则、调度计算、数据管理与消息推送；前端负责多角色交互与可视化展示；数据库负责业务数据持久化；缓存用于高频状态管理。')
add_para(doc, '实现方法上，先搭建主流程，再逐步增强：第一阶段完成订单主链路跑通；第二阶段完成骑手端与地图节点模型；第三阶段完成调度策略优化与联调；第四阶段完成测试、指标统计与文档整理。')
add_para(doc, '在调度实现方面，采用“规则筛选 + 贪婪选择”的工程策略：先根据骑手可用状态、距离与负载进行候选过滤，再基于近邻优先思想生成配送顺序，以兼顾实时性与实现复杂度。')

add_h1(doc, '五、课题的特色与预期成果')
add_para(doc, '课题特色：')
add_para(doc, '（1）以真实校园场景为约束，强调工程落地与业务闭环。')
add_para(doc, '（2）覆盖用户端、骑手端、管理端三端协同，突出系统完整性。')
add_para(doc, '（3）引入节点地图与路径可视化，提升调度过程透明度。')
add_para(doc, '（4）调度策略强调可实现、可联调、可迭代。')
add_para(doc, '预期成果：')
add_para(doc, '（1）形成一套可运行的校园外卖实时调度系统。')
add_para(doc, '（2）完成调度与路径推荐核心模块并实现可视化展示。')
add_para(doc, '（3）完成联调测试与场景测试，形成测试记录。')
add_para(doc, '（4）完成毕业设计论文及答辩演示材料。')

add_h1(doc, '六、进度安排')

schedule_rows = [
    ('01', '2025.12.29–2026.01.02\n2026.01.05–2026.01.09', '查阅校园配送与调度相关资料，明确系统目标、角色与核心流程；完成选题论证材料。', '01\n02 完成选题环节'),
    ('02', '2026.01.12–2026.01.16\n2026.01.19–2026.01.23', '对任务书进行分解，梳理功能清单与技术路线；输出模块划分和初版实施计划。', '03\n04 完成任务书环节'),
    ('03', '2026.01.26–2026.03.06\n2026.03.09–2026.03.13', '完成开题报告；同步完成系统总体设计（架构、数据库、接口主链路、测试方案）。', '05 提交开题报告，学院组织检查'),
    ('04', '2026.03.16–2026.03.20\n2026.03.23–2026.03.27\n2026.03.30–2026.04.03\n2026.04.06–2026.04.10', '完成程序开发与联调：\n1) 用户端下单与订单跟踪；\n2) 骑手端接单、状态流转与配送反馈；\n3) 管理端订单调度与运营看板；\n4) 校园节点地图、路径高亮与实时状态更新；\n5) 调度规则与贪婪路径推荐实现；\n6) 完成功能测试、接口联调与演示脚本。', '06\n07\n08\n09 提交毕设程序并进行程序演示，学院组织检查'),
    ('05', '2026.04.13–2026.04.17\n2026.04.20–2026.04.24\n2026.04.27–2026.05.01\n2026.05.04–2026.05.08', '撰写毕业论文：完成需求分析、系统设计、关键实现、联调测试与结果分析等章节。', '10\n11\n12\n13 完成并提交毕业设计（论文）'),
    ('06', '2026.05.11–2026.05.15\n2026.05.18–2026.05.22', '完成论文查重、修改完善与评阅材料准备，按意见进行针对性修订。', '14\n15 完成论文查重检测、修改及评阅'),
    ('07', '2026.05.25–2026.05.29', '完成答辩 PPT、系统演示环境与问题清单准备，参加毕业论文答辩。', '16 答辩资格审核，学院组织答辩'),
]

table = doc.add_table(rows=1, cols=4)
header = table.rows[0].cells
header[0].text = '事项'
header[1].text = '日期时间'
header[2].text = '主要工作内容'
header[3].text = '周次及时间节点'

for row in schedule_rows:
    cells = table.add_row().cells
    for i, v in enumerate(row):
        cells[i].text = v

# 表头加粗居中
for idx, cell in enumerate(table.rows[0].cells):
    for p in cell.paragraphs:
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in p.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.size = Pt(12)

# 正文字体
for r in table.rows[1:]:
    for c in r.cells:
        for p in c.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(11)

set_table_three_line(table)

add_h1(doc, '七、参考文献')
refs = [
    '[1] Cano J, Cortés P, Campo E, Correa-Espinal A A. Multi-objective grouping genetic algorithm for the joint order batching, batch assignment, and sequencing problem[J]. International Journal of Management Science and Engineering Management, 2021.',
    '[2] Simoni M D, Winkenbach M. Crowdsourced on-demand food delivery: An order batching and assignment algorithm[J]. Transportation Research Part C: Emerging Technologies, 2023.',
    '[3] Li J, Yang S, Pan W, Xu Z, Wei B. Meal delivery routing optimization with order allocation strategy based on transfer stations for instant logistics services[J]. IET Intelligent Transport Systems, 2022.',
    '[4] 李章恒. 校园外卖系统设计与实现[D]. 山东大学, 2022.',
    '[5] 杨栋博, 严张凌. 基于Spring Cloud微服务架构的校园外卖应用开发[J]. 信息与电脑(理论版), 2020, 32(13):56-58.',
    '[6] 胡秀华, 宋艳妮, 王长元. 基于移动平台的点餐系统设计与实现[J]. 电子技术与软件工程, 2018(15):39-40.'
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.5

out = Path(r'D:\new1\hanye-take-out\开题报告_重写版_三线表.docx')
doc.save(out)
print(str(out))
