# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from pathlib import Path


def add_title(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    p.paragraph_format.alignment = 1


def add_h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)


def add_body(doc, text):
    for blk in text.split('\n\n'):
        doc.add_paragraph(blk)


doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.font.size = Pt(12)

add_title(doc, '校园外卖实时调度系统开题报告（工程实践类）')
doc.add_paragraph('')

add_h2(doc, '一、系统要做什么')
add_body(doc, '本课题拟开发一套校园外卖实时调度系统，覆盖“用户下单—订单派发—骑手配送—管理监控”全流程，系统定位为工程实践项目。\n\n系统包含三个业务角色：\n1. 用户端：下单、查看订单状态、查看配送进度；\n2. 骑手端：接单、查看推荐配送顺序、更新配送状态；\n3. 管理端：订单管理、骑手管理、运营看板、异常处理。\n\n系统后端提供统一接口服务，前端提供交互页面，并通过地图可视化展示订单与骑手位置，实现调度过程可见、可追踪、可回放。')

add_h2(doc, '二、为什么要做')
add_body(doc, '校园配送场景与普通外卖场景不同，具有“点位集中、峰值并发高、骑手数量有限、时效要求高”的特点。实际运行中常见问题包括：\n1. 高峰期订单集中，人工派单效率低；\n2. 骑手路径不合理，存在绕路与等待；\n3. 配送过程不透明，管理端难以实时掌握状态；\n4. 异常订单缺少统一处理流程。\n\n因此需要建设一套可实际运行的调度系统，以工程化手段提升派单效率、配送效率和管理效率。')

add_h2(doc, '三、准备怎么做（总体方案）')
add_body(doc, '本课题采用“前后端分离 + 业务规则调度 + 地图可视化 + 联调测试”实施路线。\n\n总体技术方案如下：\n1. 前端：Web 管理端 + 移动端页面，完成下单、接单、状态流转和可视化交互；\n2. 后端：提供订单、用户、骑手、调度、统计等接口；\n3. 数据层：设计订单、骑手、节点、边、调度任务等核心数据表；\n4. 调度层：实现规则派单与贪婪路径推荐；\n5. 通知层：实时推送来单、催单、状态变更消息；\n6. 测试层：进行功能、联调、并发场景和稳定性测试。\n\n项目以“先跑通主链路，再逐步增强”的方式推进：先完成下单到配送完成，再优化调度策略和可视化效果。')

add_h2(doc, '四、会做哪些模块（工程模块划分）')
add_body(doc, '（一）用户端模块\n- 用户登录与身份识别\n- 菜品浏览与下单\n- 订单状态查询与轨迹查看\n\n（二）骑手端模块\n- 骑手登录与在线状态\n- 接单与任务列表\n- 推荐配送顺序展示\n- 配送过程状态上报\n\n（三）管理端模块\n- 订单全流程管理（待接单、配送中、完成、取消）\n- 骑手信息与负载监控\n- 校园节点地图总览\n- 异常订单处理与运营统计\n\n（四）后端服务模块\n- 订单服务、用户服务、骑手服务\n- 调度服务（派单 + 路径推荐）\n- 实时消息推送服务\n- 定时任务与状态维护\n\n（五）调度与路径模块\n- 规则派单：距离优先、负载均衡、状态约束\n- 路径推荐：基于贪婪思想生成配送顺序\n- 指标统计：平均配送时长、超时率、骑手负载分布\n\n（六）地图与可视化模块\n- 校园节点与道路边建模\n- 订单点、骑手点、路径高亮展示\n- 配送过程动态更新')

add_h2(doc, '五、联调与测试安排')
add_body(doc, '本课题将重点做工程联调与可运行性验证，而非理论证明。测试与验证方式包括：\n1. 功能测试：逐模块验证下单、接单、派单、配送、完成流程；\n2. 联调测试：前端接口调用、后端处理、数据库落库全链路验证；\n3. 场景测试：模拟饭点订单集中场景，验证系统稳定性；\n4. 对比测试：对比优化前后在配送时长、超时率方面的变化；\n5. 演示测试：准备完整演示脚本，支持学院检查与答辩展示。')

add_h2(doc, '六、最后能做到什么程度（预期成果）')
add_body(doc, '预期达到以下工程完成度：\n1. 完成可运行系统，支持用户端、骑手端、管理端基本业务闭环；\n2. 完成校园节点地图与调度可视化，支持路径与状态展示；\n3. 完成规则派单与路径推荐实现，具备基础优化效果；\n4. 完成核心联调测试与场景测试，形成测试记录与结果分析；\n5. 完成毕业设计论文与演示材料，满足工程实践类毕业设计要求。\n\n本课题不以提出全新学术模型为目标，重点在于把系统“设计出来、实现出来、联调跑通、测试验证、可展示可复现”。')

add_h2(doc, '七、进度安排（按学院时间节点）')

schedule = [
('01', '2025.12.29–2026.01.02；2026.01.05–2026.01.09', '查阅资料、完成选题', '01–02（完成选题环节）'),
('02', '2026.01.12–2026.01.16；2026.01.19–2026.01.23', '学习任务书，细化需求', '03–04（完成任务书环节）'),
('03', '2026.01.26–2026.03.06；2026.03.09–2026.03.13', '完成开题报告与方案设计', '05（提交开题报告，学院检查）'),
('04', '2026.03.16–2026.03.20；2026.03.23–2026.03.27；2026.03.30–2026.04.03；2026.04.06–2026.04.10', '系统开发、联调、阶段演示', '06–09（提交毕设程序并演示）'),
('05', '2026.04.13–2026.04.17；2026.04.20–2026.04.24；2026.04.27–2026.05.01；2026.05.04–2026.05.08', '撰写论文与整理测试结果', '10–13（完成并提交论文）'),
('06', '2026.05.11–2026.05.15；2026.05.18–2026.05.22', '查重、修改、评阅', '14–15（完成修订）'),
('07', '2026.05.25–2026.05.29', '答辩准备与正式答辩', '16（学院组织答辩）'),
]

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
h = table.rows[0].cells
h[0].text = '阶段'
h[1].text = '日期时间'
h[2].text = '主要工作内容'
h[3].text = '周次及节点'
for row in schedule:
    c = table.add_row().cells
    for i, v in enumerate(row):
        c[i].text = v

add_h2(doc, '八、参考资料')
refs = [
'[1] Cano J, Cortés P, Campo E, Correa-Espinal A A. Multi-objective grouping genetic algorithm for the joint order batching, batch assignment, and sequencing problem[J]. International Journal of Management Science and Engineering Management, 2021.',
'[2] Simoni M D, Winkenbach M. Crowdsourced on-demand food delivery: An order batching and assignment algorithm[J]. Transportation Research Part C, 2023.',
'[3] Li J, Yang S, Pan W, Xu Z, Wei B. Meal delivery routing optimization with order allocation strategy based on transfer stations for instant logistics services[J]. IET ITS, 2022.',
'[4] 李章恒. 校园外卖系统设计与实现[D]. 山东大学, 2022.',
'[5] 杨栋博, 严张凌. 基于Spring Cloud微服务架构的校园外卖应用开发[J]. 信息与电脑(理论版), 2020.',
'[6] 胡秀华, 宋艳妮, 王长元. 基于移动平台的点餐系统设计与实现[J]. 电子技术与软件工程, 2018.'
]
for it in refs:
    doc.add_paragraph(it)

out = Path(r'D:\new1\hanye-take-out\开题报告_校园外卖实时调度系统_工程实践版.docx')
doc.save(out)
print(out)
