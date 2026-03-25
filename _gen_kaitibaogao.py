# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from pathlib import Path

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.font.size = Pt(12)

p = doc.add_paragraph()
r = p.add_run('基于规则与贪婪算法的校园外卖实时调度系统开题报告')
r.bold = True
r.font.size = Pt(16)
p.paragraph_format.alignment = 1

doc.add_paragraph('')

sections = [
('一、目的与意义',
'本课题题目为“基于规则与贪婪算法的校园外卖实时调度系统”。课题面向校园封闭管理下的外卖配送场景，聚焦“最后一公里”配送效率问题，依托现有寒夜外卖项目（D:\\new1\\hanye-take-out）进行工程化改造与扩展。项目定位为工程实践，不以理论证明为核心，而以“可运行系统 + 可验证效果 + 可落地改进”为核心目标。\n\n本课题的主要目的如下：\n1. 在现有寒夜项目基础上，完善“用户端—管理端—骑手端”三端协同流程。\n2. 引入校园节点地图模型，模拟校园路网节点与边关系，实现可视化调度与路径跟踪。\n3. 在派单与路径规划模块中，结合规则策略与贪婪思想进行算法改造，提升配送效率与稳定性。\n4. 形成一套可复用的校园外卖调度工程实现方案，为后续同类项目提供参考。\n\n本课题的实践意义在于：\n1. 提高高峰时段订单分配效率，缓解人工派单压力。\n2. 降低骑手无效绕行，缩短平均配送时长。\n3. 用可视化方式展示“订单—骑手—路线”关系，提升系统可解释性。\n4. 形成完整的软件工程成果（系统、文档、测试与分析），满足本科毕业设计工程类要求。'),
('二、研究概况（工程现状）',
'校园外卖系统近年的实现路径主要分为两类：\n1. 以业务功能为主的管理系统（订单、菜品、用户管理）；\n2. 以调度优化为主的配送系统（派单、路径规划、实时追踪）。\n\n当前多数项目在业务流程方面较完善，但在“校园特定路网 + 多订单并发 + 骑手负载均衡”的综合场景下，调度能力仍有提升空间。尤其在饭点高并发条件下，仅按最近距离分配往往会出现局部最优、全局效率不足的问题。\n\n因此，本课题拟在已有外卖业务系统之上，重点强化调度侧能力：\n1. 基于规则的自动派单（距离、负载、状态等）。\n2. 基于贪婪思想的路径推荐（最近邻/顺路优先）。\n3. 基于节点地图的路线可视化与轨迹跟踪。'),
('三、理论依据、研究内容与研究方法',
'（一）理论依据\n1. 分层架构思想：前后端分离，后端采用 Controller/Service/Mapper 分层。\n2. RESTful 接口规范：统一前后端通信方式。\n3. 身份认证与权限控制：JWT + 拦截器保障接口访问安全。\n4. 缓存与状态管理：Redis 管理店铺状态等高频状态数据。\n5. 图模型与贪婪策略：将校园道路抽象为节点与边，结合规则评分和贪婪选择实现调度。\n\n（二）研究内容（工程实现）\n1. 三端业务功能完善：用户端、管理端、骑手端流程打通。\n2. 校园路网建模：维护校园节点与边数据，支持地图可视化。\n3. 自动派单模块：结合“距离优先 + 负载均衡 + 可用状态”进行骑手选择。\n4. 路径推荐模块：在多订单情况下采用贪婪策略生成配送顺序。\n5. 实时追踪与提醒：通过 WebSocket 推送来单、催单与配送状态变化。\n6. 系统测试与效果分析：围绕效率指标开展对比验证。\n\n（三）研究方法\n1. 文献与方案调研法：梳理校园配送与路径优化相关实现思路。\n2. 原型迭代法：在寒夜项目已有代码基础上逐步扩展功能。\n3. 对比实验法：对比“改造前/改造后”在订单分配与配送时效上的差异。\n4. 场景模拟法：模拟校园节点、订单分布与骑手状态，进行压力与稳定性测试。'),
('四、课题创新点（工程创新）',
'1. 在现有外卖系统中新增骑手端，实现三端协同闭环。\n2. 引入校园节点地图，将配送过程从“黑盒派单”升级为“可视化调度”。\n3. 将规则策略与贪婪算法结合，兼顾实时性与工程可实现性。\n4. 在管理端与用户端同时提供轨迹/调度结果展示，提高系统可解释性。\n5. 形成“业务系统 + 调度算法 + 可视化地图 + 实时推送”的工程实践一体化方案。'),
('五、预期成果',
'1. 一套可运行的校园外卖调度系统（用户端、骑手端、管理端、后端服务）。\n2. 完整的数据库结构与核心业务接口。\n3. 派单与路径推荐算法实现及实验对比结果。\n4. 开题报告、中期材料、毕业论文与答辩材料。'),
]

for title, body in sections:
    h = doc.add_paragraph()
    rr = h.add_run(title)
    rr.bold = True
    rr.font.size = Pt(14)
    for para in body.split('\n\n'):
        doc.add_paragraph(para)

a = doc.add_paragraph()
r = a.add_run('六、进度安排')
r.bold = True
r.font.size = Pt(14)

schedule = [
('01', '2025.12.29–2026.01.02；2026.01.05–2026.01.09', '查阅文献、提交选题', '01–02（完成选题环节）'),
('02', '2026.01.12–2026.01.16；2026.01.19–2026.01.23', '下达并学习任务书', '03–04（完成任务书环节）'),
('03', '2026.01.26–2026.03.06；2026.03.09–2026.03.13', '文献整理、完成开题报告', '05（提交开题报告，学院检查）'),
('04', '2026.03.16–2026.03.20；2026.03.23–2026.03.27；2026.03.30–2026.04.03；2026.04.06–2026.04.10', '毕业设计程序开发（骑手端、节点地图、调度算法改造、联调测试）', '06–09（提交毕设程序并演示，学院检查）'),
('05', '2026.04.13–2026.04.17；2026.04.20–2026.04.24；2026.04.27–2026.05.01；2026.05.04–2026.05.08', '撰写毕业论文', '10–13（完成并提交毕业设计（论文））'),
('06', '2026.05.11–2026.05.15；2026.05.18–2026.05.22', '论文查重、修改、评阅', '14–15（按学院要求完成查重与修订）'),
('07', '2026.05.25–2026.05.29', '论文答辩准备与答辩', '16（答辩资格审核，学院组织答辩）'),
]

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '阶段'
hdr[1].text = '日期时间'
hdr[2].text = '主要工作内容'
hdr[3].text = '周次及节点'
for row in schedule:
    cells = table.add_row().cells
    for i,v in enumerate(row):
        cells[i].text = v

b = doc.add_paragraph()
rb = b.add_run('七、参考文献')
rb.bold = True
rb.font.size = Pt(14)

refs = [
'[1] Cano J, Cortés P, Campo E, Correa-Espinal A A. Multi-objective grouping genetic algorithm for the joint order batching, batch assignment, and sequencing problem[J]. International Journal of Management Science and Engineering Management, 2021.',
'[2] Simoni M D, Winkenbach M. Crowdsourced on-demand food delivery: An order batching and assignment algorithm[J]. Transportation Research Part C, 2023.',
'[3] Li J, Yang S, Pan W, Xu Z, Wei B. Meal delivery routing optimization with order allocation strategy based on transfer stations for instant logistics services[J]. IET ITS, 2022.',
'[4] 李章恒. 校园外卖系统设计与实现[D]. 山东大学, 2022.',
'[5] 杨栋博, 严张凌. 基于Spring Cloud微服务架构的校园外卖应用开发[J]. 信息与电脑(理论版), 2020.',
'[6] 胡秀华, 宋艳妮, 王长元. 基于移动平台的点餐系统设计与实现[J]. 电子技术与软件工程, 2018.'
]
for i in refs:
    doc.add_paragraph(i)

out = Path(r'D:\new1\hanye-take-out\开题报告_基于规则与贪婪算法的校园外卖实时调度系统.docx')
doc.save(out)
print(str(out))
